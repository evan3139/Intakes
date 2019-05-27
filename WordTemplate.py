import xlrd
from docx import Document


def create_docx_template(excel_name, docx_name, file_title):
    # Here will create a template in a docx file for the quiz templates
    workbook = xlrd.open_workbook(excel_name)
    worksheet = workbook.sheet_by_index(0)
    date = worksheet.cell_value(0, 0)

    # Creates the needed arrays
    names = []
    ages = []
    genders = []
    races = []

    # Grabs all the names ages genders and races through their columns and puts them in arrays.
    for i in range(worksheet.nrows):
        if i > 0:
            names.append(worksheet.cell_value(i, 1))
            ages.append(int(worksheet.cell_value(i, 5)))
            genders.append(worksheet.cell_value(i, 6))
            races.append(worksheet.cell_value(i, 12))
    doc = Document()

    # Creates the top of the file with stuff that will not change.
    doc.add_paragraph(date + ":")
    doc.add_paragraph("Facilitator:")
    doc.add_paragraph("Topic:")
    doc.add_paragraph("Week:")
    doc.add_paragraph("Group:" + file_title)
    doc.add_paragraph("Questions:")
    doc.add_paragraph("")

    # Fills the file with ID's genders, ages, races in a template format.
    for index, i in enumerate(names):
        doc.add_paragraph(
            str(genders[index]) + ',' + str(ages[index]) + "," + str(races[index]) + "," + str(names[index]) + ":")
    doc.save(docx_name)