from tkinter import Tk
from tkinter.filedialog import askopenfilename
import os
import xlsxwriter
from docx import Document
import openpyxl
from shutil import copyfile
from IsNumber import is_number
from ResizeColumn import resize_columns
from quiz_combined import *


# The function call for putting everything into an excel sheet.
def excel(names, numbers, headers, contents, genders, ages, races, excel_name, divide_by):
    # Do all the writing to the file here for the Heading which will never change
    workbook = xlsxwriter.Workbook(excel_name)
    worksheet = workbook.add_worksheet()

    # This is so we know what index to start writing to after the header is in place in the sheet. Couldn't think of
    # a better way
    counter = 0
    average = 0
    total = 0
    # Score after its shifted to a score out of 100
    formatted_score = 0

    for index, x in enumerate(headers):
        worksheet.write(index, 0, x)
        if index != len(contents):
            worksheet.write(index, 1, contents[index])
        counter += 1

    # Special Case as this is its own var and not in an array
    worksheet.write(counter - 1, 1, divide_by)
    counter += 1

    # This is the header in the excel sheet for Gender Age race etc.
    worksheet.write(counter, 0, "Gender")
    worksheet.write(counter, 1, "Age")
    worksheet.write(counter, 2, "Race")
    worksheet.write(counter, 3, "ID")
    worksheet.write(counter, 4, "Score")
    counter += 1

    for index, x in enumerate(names):
        if is_number(numbers[index]):
            formatted_score = numbers[index] / divide_by * 100
            average += round(formatted_score)
            worksheet.write(counter, 4, round(formatted_score))
            total += 1
        else:
            worksheet.write(counter, 4, numbers[index])
        worksheet.write(counter, 0, genders[index])
        worksheet.write(counter, 1, ages[index])
        worksheet.write(counter, 2, races[index])
        worksheet.write(counter, 3, x)
        counter += 1

    # Calculate the average and put it at the end of the Excel Sheet
    average = average / total
    worksheet.write(counter + 1, 3, "Average")
    worksheet.write(counter + 1, 4, average)

    # Save in the folder

    workbook.close()


# This will delete the template and remake it so that people dont have to delete or copy anything. since they are old they struggle with this.
def remake_template(names, headers, contents, genders, ages, races, filename, group):
    # First delete the old file that she filled out.
    os.remove(filename)
    doc = Document()
    for i in headers:
        if "group" in i.lower():
            doc.add_paragraph(i + ":" + group)
        else:
            doc.add_paragraph(i + ":")
    doc.add_paragraph("")

    for index, i in enumerate(names):
        doc.add_paragraph(genders[index] + ',' + str(ages[index]) + ',' + races[index] + ',' + names[index] + ':')
    doc.save(filename)


def read_doc(filename):
    doc = Document(filename)

    # This will create an array for scores then fulllist which will seperate names from scores
    lines = []
    full_list = []
    numbers = []
    names = []
    headers = []
    contents = []
    genders = []
    ages = []
    races = []
    divide_by = None
    group = " "

    # Reads into the scores array through docx.
    for para in doc.paragraphs:
        # If its an empty line do this
        if para.text == "":
            continue
        elif para.text.isspace():
            continue
        else:
            lines.append(para.text)

    # Creates a seperate value for Non name values and numbers etc
    for index, x in enumerate(lines):
        if "date" in x.lower():
            header, content = x.split(":")
            content = content.strip()
            headers.append(header)
            contents.append(content)
        elif "facilitator" in x.lower():
            header, content = x.split(":")
            content = content.strip()
            headers.append(header)
            contents.append(content)
        elif "topic" in x.lower():
            header, content = x.split(":")
            content = content.strip()
            headers.append(header)
            contents.append(content)
        elif "week" in x.lower():
            header, content = x.split(":")
            content = content.strip()
            headers.append(header)
            if is_number(content):
                contents.append(int(content))
            else:
                contents.append(content)
        elif "group" in x.lower():
            header, content = x.split(":")
            # we need to keep the group, it is important for remaking the file.
            group = content
            content = content.strip()
            headers.append(header)
            contents.append(content)
            # Gotta change the var to divide by since its total questions asked per person
        elif "questions" in x.lower():
            header, div = x.split(":")
            headers.append(header)
            div = div.strip()
            try:
                divide_by = int(div)
            except ValueError:
                raise Exception("Please enter a valid number (Digit not word) For Questions")
        else:
            name, score = x.split(":")
            if is_number(score):
                # Do the math for right vs total
                score_num = int(score) / divide_by
                numbers.append(int(score))
            else:
                score = score.strip()
                numbers.append(score)
            gender, age, race, ID = name.split(",")
            genders.append(gender)
            ages.append(int(age))
            races.append(race)
            names.append(ID)

    newpath = 'C:/VantagePoint/Quizzes' + "/" + str(contents[4].strip())
    if not os.path.exists(newpath):
        os.makedirs(newpath)

    excel_name = os.path.join(newpath + "/" + str(contents[2]) + "-" + str(contents[4]) + ".xlsx")
    docx_name = os.path.join(newpath + "/" + str(contents[2]) + "-" + str(contents[4]) + ".docx")

    # Copy the file so we can reset the original to make it more user friendly
    if not filename == docx_name:
        copyfile(filename, docx_name)

    # remove all whitespace just in case
    names = [x.strip(' ') for x in names]

    # Call all the functions
    excel(names, numbers, headers, contents, genders, ages, races, excel_name, divide_by)
    resize_columns(excel_name)
    if not filename == docx_name:
        remake_template(names, headers, contents, genders, ages, races, filename, group)
